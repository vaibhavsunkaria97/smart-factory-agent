"""Generates the technical report PDF and DOCX from real pipeline results."""
import json
import os
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Image,
                                Table, TableStyle, HRFlowable, KeepTogether)
from PIL import Image as PILImage
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ---- EDIT THESE ----------------------------------------------------------
AUTHOR = "Vaibhav Kumar Sunkaria"                      # <-- your full name
GITHUB = "https://github.com/vaibhavsunkaria97/smart-factory-agent"
# --------------------------------------------------------------------------

DARK = colors.HexColor("#0f172a")
ACCENT = colors.HexColor("#1d4ed8")
SLATE = colors.HexColor("#475569")
LIGHT = colors.HexColor("#f1f5f9")

S = getSampleStyleSheet()
S.add(ParagraphStyle("Body2", parent=S["Normal"], fontSize=8.6, leading=11.5,
                     alignment=TA_JUSTIFY, spaceAfter=4))
S.add(ParagraphStyle("H1c", parent=S["Heading1"], fontSize=12, textColor=DARK,
                     spaceBefore=7, spaceAfter=2))
S.add(ParagraphStyle("H2c", parent=S["Heading2"], fontSize=9.6,
                     textColor=ACCENT, spaceBefore=5, spaceAfter=1))
S.add(ParagraphStyle("Cap", parent=S["Normal"], fontSize=7.2, textColor=SLATE,
                     alignment=TA_CENTER, spaceAfter=6))
S.add(ParagraphStyle("TitleBig", parent=S["Title"], fontSize=17,
                     textColor=DARK, spaceAfter=1))
S.add(ParagraphStyle("Sub", parent=S["Normal"], fontSize=9, textColor=SLATE,
                     alignment=TA_CENTER, spaceAfter=1))
cell = ParagraphStyle("cell", parent=S["Normal"], fontSize=7.6, leading=9.6)
hcell = ParagraphStyle("hcell", parent=cell, textColor=colors.white,
                       fontName="Helvetica-Bold")

CW = letter[0] - 1.4 * inch
ev = json.load(open("evaluation_results.json"))
res = ev["results"]


def img(path, max_w=CW, max_h=None):
    w, h = PILImage.open(path).size
    s = max_w / w
    if max_h and h * s > max_h:
        s = max_h / h
    im = Image(path, width=w * s, height=h * s)
    im.hAlign = 'CENTER'
    return im


def C(t, st=cell): return Paragraph(t, st)
def h1(n, t): return Paragraph(f"{n}&nbsp;&nbsp;{t}", S["H1c"])
def h2(t): return Paragraph(t, S["H2c"])
def p(t): return Paragraph(t, S["Body2"])


def styled(tbl, head=DARK):
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), head),
        ("FONTSIZE", (0, 0), (-1, -1), 7.6),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, LIGHT]),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cbd5e1")),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
    ]))
    tbl.hAlign = 'CENTER'
    return tbl


story = [
    Paragraph("Smart Factory Equipment Anomaly Alert Agent", S["TitleBig"]),
    Paragraph("Pegatron ML Engineer Assignment — Assignment 3", S["Sub"]),
    Paragraph(f"{AUTHOR}&nbsp;&nbsp;·&nbsp;&nbsp;Built end-to-end with free AI tools"
              f"&nbsp;&nbsp;·&nbsp;&nbsp;<font face='Courier'>{GITHUB}</font>", S["Sub"]),
    Spacer(1, 4), HRFlowable(width="100%", thickness=1.1, color=DARK), Spacer(1, 4),

    h1("1", "Task Description &amp; Objective"),
    p("An AI agent that monitors smart-factory equipment sensor data, detects anomalies "
      "automatically, and delivers <b>actionable real-time alerts</b>. The pipeline generates a "
      "dummy sensor dataset (temperature, pressure, vibration), cleans it, scores every reading "
      "with two independent detectors, explains each anomaly through an LLM, and emits a "
      "prioritised alert table plus a machine-readable JSON log for downstream MES integration."),
    p("<b>Design constraint taken from the brief.</b> Because the assignment rewards a low usage "
      "barrier, the system runs with <b>one command and no API key</b>: "
      "<font face='Courier'>python run_all.py</font> regenerates the data, alerts, metrics, "
      "figures and test results. LLM reasoning is used when available and degrades gracefully "
      "to a deterministic expert reasoner when it is not."),

    h1("2", "Development Workflow (AI-driven)"),
    p("The project was built with an <b>agentic coding loop</b>: Aider driving a free-tier "
      "frontier model, reading the repository, writing files and committing to git. Manual effort "
      "was limited to specifying constraints, verifying every output by execution, and taking "
      "over where the agent stalled. Full prompts and the defect log are in "
      "<font face='Courier'>docs/AI_LOG.md</font>."),

    h1("3", "Method &amp; Model Details"),
    h2("3.1  Data generation and preprocessing"),
    p("300 rows at one-minute intervals, ~11% abnormal, with 1–3 channels corrupted per abnormal "
      "row. Missing values and duplicate timestamps are injected deliberately so cleaning is not a "
      "no-op. A generator-side assertion guarantees every row labelled abnormal breaches at least "
      "one threshold. Preprocessing de-duplicates, imputes with time-aware interpolation, and "
      "produces a <b>separate z-scored copy</b> for the ML detector — physical units are kept "
      "apart from scaled values so alerts stay human-readable."),
    h2("3.2  Detection ensemble"),
    p("A <b>rule detector</b> applies the brief's physical thresholds (temp &gt;52/&lt;43 °C, "
      "pressure &gt;1.08/&lt;0.97 bar, vibration &gt;0.07 g) — fast, explainable, no training. An "
      "<b>IsolationForest</b> (200 trees) learns the normal multivariate envelope unsupervised and "
      "catches patterns no single threshold expresses. Scores are fused as an OR-of-evidence "
      "(<font face='Courier'>score = max(rule, ML)</font>), favouring recall because a missed "
      "fault costs more than a false alarm. Severity follows from breach count and score."),
    h2("3.3  Reasoning layer"),
    p("Each anomaly is turned into a diagnosis and recommended action by the best available free "
      "backend — local <b>Ollama</b>, free-tier <b>Groq</b>, or a deterministic knowledge base "
      "keyed on (signal, direction). The LLM never alters a detection; it only phrases the "
      "explanation, so correctness remains deterministic and unit-tested."),
]

story.append(KeepTogether([
    h1("4", "System Architecture"),
    img("assets/architecture.png", max_w=CW * 0.90, max_h=4.2 * inch),
    Paragraph("Figure 1 — Four-stage pipeline with graceful degradation at the reasoning layer.",
              S["Cap"]),
]))

order = [k for k in res if "shipped" not in k] + [k for k in res if "shipped" in k]
rows = [[C("Detector", hcell), C("Precision", hcell), C("Recall", hcell),
         C("F1", hcell), C("F2", hcell), C("Flagged", hcell)]]
for k in order:
    m = res[k]
    b, e = ("<b>", "</b>") if "shipped" in k else ("", "")
    rows.append([C(f"{b}{k}{e}"), C(f"{b}{m['precision']:.3f}{e}"),
                 C(f"{b}{m['recall']:.3f}{e}"), C(f"{b}{m['f1']:.3f}{e}"),
                 C(f"{b}{m['f2']:.3f}{e}"), C(f"{b}{m['flagged']}{e}")])

story += [
    h1("5", "Evaluation"),
    p(f"Four configurations were compared against ground-truth labels "
      f"({ev['n_rows']} readings, {ev['n_true_anomalies']} true anomalies). F2 is reported "
      f"alongside F1 because β=2 weights recall four times more than precision, which is the "
      f"right emphasis for equipment safety. Accuracy is deliberately omitted: with ~11% "
      f"positives, predicting \"all normal\" would already score ~89%."),
    styled(Table(rows, colWidths=[2.25*inch, 0.78*inch, 0.68*inch, 0.6*inch,
                                  0.6*inch, 0.7*inch])),
    Paragraph("Table 1 — Detector comparison. All configurations achieve perfect recall.", S["Cap"]),
    p("<b>An honest finding: preprocessing manufactured two anomalies.</b> Precision is below 1.0 "
      "for every detector, and investigation showed why. Two rows with a missing vibration reading "
      "sat between high-vibration neighbours; time-interpolation filled them with 0.080 and 0.088, "
      "both above the 0.07 limit. Those rows are labelled normal but breach after cleaning, so "
      "they count as false positives. The detectors behaved correctly — the imputed value never "
      "occurred. <b>Cleaning is not a neutral step.</b> In production the fix is causal "
      "forward-fill (which the streaming path requires anyway) plus flagging imputed values in the "
      "alert so an operator knows the reading was inferred rather than measured."),
]

story.append(KeepTogether([
    img("assets/evaluation.png", max_w=CW, max_h=2.8 * inch),
    Paragraph("Figure 2 — Confusion matrix and precision–recall curve with the shipped "
              "operating point.", S["Cap"]),
]))

story.append(KeepTogether([
    h1("6", "Demo"),
    img("assets/demo_llm.png", max_w=CW, max_h=3.6 * inch),
    Paragraph("Figure 3 — CRITICAL alerts with live Ollama reasoning: cleaning summary, "
              "prioritised table, and the top-priority alert card.", S["Cap"]),
    Spacer(1, 6),
    img("assets/demo_offline.png", max_w=CW, max_h=3.6 * inch),
    Paragraph("Figure 4 — Offline fallback execution with deterministic expert reasoner output "
              "when no LLM backend is active.", S["Cap"]),
]))

ai = [
    [C("Free AI tool", hcell), C("Type", hcell), C("Contribution", hcell)],
    [C("<b>Aider</b> + free-tier frontier model"), C("Build-time"),
     C("Agentic loop: read repo, planned changes, wrote modules and tests, committed to git.")],
    [C("<b>Claude</b> (free tier)"), C("Build-time"),
     C("Requirement analysis, architecture decisions, debugging, report generation.")],
    [C("<b>Ollama</b> / llama3.2 (local)"), C("Runtime"),
     C("LLM inside the product writing each alert's diagnosis — free, offline, air-gap friendly.")],
    [C("<b>Groq</b> (free tier)"), C("Runtime"),
     C("Hosted alternative reasoning backend.")],
]

story += [
    h1("7", "Use of AI Tools"),
    p("Only free tiers were used. AI appears in two distinct roles: tools that <b>built</b> the "
      "system, and an LLM embedded <b>inside</b> the delivered product as its reasoning layer."),
    styled(Table(ai, colWidths=[1.75*inch, 0.75*inch, 4.1*inch]), head=ACCENT),
    Spacer(1, 3),
    h2("7.1  AI as reviewer, not only generator"),
    p("Every AI output was verified by execution, and verification caught defects that reading the "
      "code did not: mislabelled rows where 8 of 41 'abnormal' readings breached no threshold; a "
      "syntax error from a duplicated loop header; a runtime TypeError from interpolating a string "
      "column; a scale-dependent bug where 100-row datasets received zero missing values; and — "
      "most instructive — an LLM backend that reported success while every call silently failed, "
      "because it posted to the wrong endpoint and a bare <font face='Courier'>except: pass</font> "
      "hid it. The status line said \"ollama\" while zero LLM calls had succeeded. Report observed "
      "state, never intended state."),
    p("<b>Measured limits of the local model.</b> On multi-breach rows llama3.2 frequently names "
      "only one fault, occasionally inverts the direction (recommending \"increase temperature\" "
      "for a low-temperature fault), and has hallucinated non-existent hardware (\"Replace sensor "
      "3\"). The deterministic knowledge base never does. This is the empirical basis for keeping "
      "detection deterministic and confining the LLM to phrasing."),

    h1("8", "Limitations &amp; Next Steps"),
    p("The system detects <b>point anomalies</b> only; contextual and collective (subsequence) "
      "faults would need windowed or sequence models — deliberately not attempted here because "
      "the synthetic data contains independent point anomalies with no temporal structure for a "
      "sequence model to learn. Batch imputation is non-causal and unsuitable for streaming. "
      "Severity cutoffs are a heuristic rather than a derivation. Next steps: Kafka ingestion in "
      "place of CSV, per-equipment learned thresholds, drift detection with retraining, an "
      "operator confirm/dismiss feedback loop to build real labels, retrieval-grounded "
      "maintenance actions, and MES integration so a CRITICAL alert opens a work order."),
]

doc = SimpleDocTemplate("Technical_Report.pdf", pagesize=letter,
                        topMargin=0.55*inch, bottomMargin=0.5*inch,
                        leftMargin=0.7*inch, rightMargin=0.7*inch,
                        title="Smart Factory Alert Agent — Technical Report",
                        author=AUTHOR)
doc.build(story)
print("wrote Technical_Report.pdf")


def build_docx():
    doc = docx.Document()
    for s in doc.sections:
        s.top_margin = Inches(0.75)
        s.bottom_margin = Inches(0.75)
        s.left_margin = Inches(0.75)
        s.right_margin = Inches(0.75)

    title = doc.add_heading("Smart Factory Equipment Anomaly Alert Agent", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in title.runs:
        r.font.size = Pt(18)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)

    sub = doc.add_paragraph("Pegatron ML Engineer Assignment — Assignment 3")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(10)
    sub.runs[0].font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    author_p = doc.add_paragraph(f"{AUTHOR}  ·  Built end-to-end with free AI tools  ·  {GITHUB}")
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_p.runs[0].font.size = Pt(9.5)
    author_p.runs[0].font.color.rgb = RGBColor(0x1d, 0x4e, 0xd8)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def add_h1(num, text):
        h = doc.add_heading(f"{num}.  {text}", level=1)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        for r in h.runs:
            r.font.size = Pt(13)
            r.font.bold = True
            r.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)

    def add_h2(text):
        h = doc.add_heading(text, level=2)
        h.paragraph_format.space_before = Pt(8)
        h.paragraph_format.space_after = Pt(2)
        for r in h.runs:
            r.font.size = Pt(11)
            r.font.bold = True
            r.font.color.rgb = RGBColor(0x1d, 0x4e, 0xd8)

    def add_p(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        import re
        parts = re.split(r'(<b>.*?</b>)', text)
        for part in parts:
            if part.startswith('<b>') and part.endswith('</b>'):
                r = p.add_run(part[3:-4])
                r.font.bold = True
            else:
                p.add_run(part)

    def add_img(path, caption, width_in=6.5):
        if os.path.exists(path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run()
            run.add_picture(path, width=Inches(width_in))
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_after = Pt(8)
            for r in cap.runs:
                r.font.size = Pt(8.5)
                r.font.italic = True
                r.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    add_h1("1", "Task Description & Objective")
    add_p("An AI agent that monitors smart-factory equipment sensor data, detects anomalies automatically, and delivers actionable real-time alerts. The pipeline generates a dummy sensor dataset (temperature, pressure, vibration), cleans it, scores every reading with two independent detectors, explains each anomaly through an LLM, and emits a prioritised alert table plus a machine-readable JSON log for downstream MES integration.")
    add_p("Design constraint taken from the brief. Because the assignment rewards a low usage barrier, the system runs with one command and no API key: python run_all.py regenerates the data, alerts, metrics, figures and test results. LLM reasoning is used when available and degrades gracefully to a deterministic expert reasoner when it is not.")

    add_h1("2", "Development Workflow (AI-driven)")
    add_p("The project was built with an agentic coding loop: Aider driving a free-tier frontier model, reading the repository, writing files and committing to git. Manual effort was limited to specifying constraints, verifying every output by execution, and taking over where the agent stalled. Full prompts and the defect log are in docs/AI_LOG.md.")

    add_h1("3", "Method & Model Details")
    add_h2("3.1  Data generation and preprocessing")
    add_p("300 rows at one-minute intervals, ~11% abnormal, with 1–3 channels corrupted per abnormal row. Missing values and duplicate timestamps are injected deliberately so cleaning is not a no-op. A generator-side assertion guarantees every row labelled abnormal breaches at least one threshold. Preprocessing de-duplicates, imputes with time-aware interpolation, and produces a separate z-scored copy for the ML detector — physical units are kept apart from scaled values so alerts stay human-readable.")

    add_h2("3.2  Detection ensemble")
    add_p("A rule detector applies the brief's physical thresholds (temp >52/<43 °C, pressure >1.08/<0.97 bar, vibration >0.07 g) — fast, explainable, no training. An IsolationForest (200 trees) learns the normal multivariate envelope unsupervised and catches patterns no single threshold expresses. Scores are fused as an OR-of-evidence (score = max(rule, ML)), favouring recall because a missed fault costs more than a false alarm. Severity follows from breach count and score.")

    add_h2("3.3  Reasoning layer")
    add_p("Each anomaly is turned into a diagnosis and recommended action by the best available free backend — local Ollama, free-tier Groq, or a deterministic knowledge base keyed on (signal, direction). The LLM never alters a detection; it only phrases the explanation, so correctness remains deterministic and unit-tested.")

    add_h1("4", "System Architecture")
    add_img("assets/architecture.png", "Figure 1 — Four-stage pipeline with graceful degradation at the reasoning layer.", width_in=6.0)

    add_h1("5", "Evaluation")
    add_p(f"Four configurations were compared against ground-truth labels ({ev['n_rows']} readings, {ev['n_true_anomalies']} true anomalies). F2 is reported alongside F1 because β=2 weights recall four times more than precision, which is the right emphasis for equipment safety. Accuracy is deliberately omitted: with ~11% positives, predicting 'all normal' would already score ~89%.")

    # Table 1 in DOCX
    t1 = doc.add_table(rows=1, cols=6)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = t1.rows[0].cells
    headers = ["Detector", "Precision", "Recall", "F1", "F2", "Flagged"]
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p_cell in hdr_cells[i].paragraphs:
            p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p_cell.runs:
                r.font.bold = True

    for k in order:
        m = res[k]
        row_cells = t1.add_row().cells
        row_cells[0].text = k
        row_cells[1].text = f"{m['precision']:.3f}"
        row_cells[2].text = f"{m['recall']:.3f}"
        row_cells[3].text = f"{m['f1']:.3f}"
        row_cells[4].text = f"{m['f2']:.3f}"
        row_cells[5].text = f"{m['flagged']}"
        for i, cell in enumerate(row_cells):
            for p_cell in cell.paragraphs:
                p_cell.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER

    cap1 = doc.add_paragraph("Table 1 — Detector comparison. All configurations achieve perfect recall.")
    cap1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap1.paragraph_format.space_after = Pt(8)
    for r in cap1.runs:
        r.font.size = Pt(8.5)
        r.font.italic = True

    add_p("An honest finding: preprocessing manufactured two anomalies. Precision is below 1.0 for every detector, and investigation showed why. Two rows with a missing vibration reading sat between high-vibration neighbours; time-interpolation filled them with 0.080 and 0.088, both above the 0.07 limit. Those rows are labelled normal but breach after cleaning, so they count as false positives. The detectors behaved correctly — the imputed value never occurred. Cleaning is not a neutral step. In production the fix is causal forward-fill (which the streaming path requires anyway) plus flagging imputed values in the alert so an operator knows the reading was inferred rather than measured.")

    add_img("assets/evaluation.png", "Figure 2 — Confusion matrix and precision–recall curve with the shipped operating point.", width_in=6.5)

    add_h1("6", "Demo")
    add_img("assets/demo_llm.png", "Figure 3 — CRITICAL alerts with live Ollama reasoning: cleaning summary, prioritised table, and top-priority alert card.", width_in=6.5)
    add_img("assets/demo_offline.png", "Figure 4 — Offline fallback execution with deterministic expert reasoner output when no LLM backend is active.", width_in=6.5)

    add_h1("7", "Use of AI Tools")
    add_p("Only free tiers were used. AI appears in two distinct roles: tools that built the system, and an LLM embedded inside the delivered product as its reasoning layer.")

    t2 = doc.add_table(rows=1, cols=3)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr2 = t2.rows[0].cells
    hdr2[0].text = "Free AI tool"
    hdr2[1].text = "Type"
    hdr2[2].text = "Contribution"
    for cell in hdr2:
        for p_cell in cell.paragraphs:
            for r in p_cell.runs:
                r.font.bold = True

    for item in [
        ("Aider + free-tier frontier model", "Build-time", "Agentic loop: read repo, planned changes, wrote modules and tests, committed to git."),
        ("Claude (free tier)", "Build-time", "Requirement analysis, architecture decisions, debugging, report generation."),
        ("Ollama / llama3.2 (local)", "Runtime", "LLM inside the product writing each alert's diagnosis — free, offline, air-gap friendly."),
        ("Groq (free tier)", "Runtime", "Hosted alternative reasoning backend.")
    ]:
        row_cells = t2.add_row().cells
        row_cells[0].text = item[0]
        row_cells[1].text = item[1]
        row_cells[2].text = item[2]

    add_h2("7.1  AI as reviewer, not only generator")
    add_p("Every AI output was verified by execution, and verification caught defects that reading the code did not: mislabelled rows where 8 of 41 'abnormal' readings breached no threshold; a syntax error from a duplicated loop header; a runtime TypeError from interpolating a string column; a scale-dependent bug where 100-row datasets received zero missing values; and — most instructive — an LLM backend that reported success while every call silently failed, because it posted to the wrong endpoint and a bare except: pass hid it. The status line said 'ollama' while zero LLM calls had succeeded. Report observed state, never intended state.")
    add_p("Measured limits of the local model. On multi-breach rows llama3.2 frequently names only one fault, occasionally inverts the direction (recommending 'increase temperature' for a low-temperature fault), and has hallucinated non-existent hardware ('Replace sensor 3'). The deterministic knowledge base never does. This is the empirical basis for keeping detection deterministic and confining the LLM to phrasing.")

    add_h1("8", "Limitations & Next Steps")
    add_p("The system detects point anomalies only; contextual and collective (subsequence) faults would need windowed or sequence models — deliberately not attempted here because the synthetic data contains independent point anomalies with no temporal structure for a sequence model to learn. Batch imputation is non-causal and unsuitable for streaming. Severity cutoffs are a heuristic rather than a derivation. Next steps: Kafka ingestion in place of CSV, per-equipment learned thresholds, drift detection with retraining, an operator confirm/dismiss feedback loop to build real labels, retrieval-grounded maintenance actions, and MES integration so a CRITICAL alert opens a work order.")

    doc.save("Technical_Report.docx")
    print("wrote Technical_Report.docx")


build_docx()